"""
docx.py — Сохранение диалогов в документы Word (DOCX)
======================================================

ПРОСТЫМИ СЛОВАМИ:
Каждый диалог с ботом и каждый голосовой запрос сохраняются
в отдельный Word-документ в папке data/docs. Это «бумажная» копия памяти.
"""

import time
from datetime import datetime

from docx import Document

from bot import config


def save_docx(user_id, user_text: str, ai_text: str) -> str:
    """
    Сохраняет один обмен репликами (вопрос пользователя + ответ бота)
    в документ Word. Возвращает путь к созданному файлу.
    """
    doc = Document()

    doc.add_heading("AI MEMORY", level=1)
    doc.add_paragraph(f"DATE: {datetime.now()}")

    doc.add_heading("USER", level=2)
    doc.add_paragraph(user_text)

    doc.add_heading("AI", level=2)
    doc.add_paragraph(ai_text)

    filename = str(config.DOCS_DIR / f"{user_id}_{int(time.time())}.docx")
    doc.save(filename)

    return filename


def save_voice_transcript(user_id, text: str) -> str:
    """
    Сохраняет распознанный текст голосового сообщения в документ Word.
    Возвращает путь к созданному файлу.
    """
    doc = Document()

    doc.add_heading("VOICE TRANSCRIPT", level=1)
    doc.add_paragraph(text)

    filename = str(config.DOCS_DIR / f"voice_{user_id}.docx")
    doc.save(filename)

    return filename
