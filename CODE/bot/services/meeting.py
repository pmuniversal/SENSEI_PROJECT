"""
meeting.py — Генерация протокола (Bayonnoma) и повестки (Agenda) по шаблону Pulinform
======================================================================================

Принимает транскрипт собрания → GPT-4.1 структурирует → возвращает два DOCX файла.

Формат файлов:
  ДД_Месяц_ГГГГ_Bayonnoma.docx — протокол прошедшего собрания
  ДД+1_Месяц_ГГГГ_Agenda.docx  — повестка следующего собрания
"""

import json
import re
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from bot import config
from bot.services.ai import client

# Цветовая кодировка строк
COLORS = {
    "red":    "F8D7DA",  # критично
    "yellow": "FFF3CD",  # внимание
    "green":  "D4EDDA",  # выполнено
    "white":  None,      # нейтрально
}

MONTHS_UZ = {
    1: "Yanvar", 2: "Fevral", 3: "Mart", 4: "Aprel",
    5: "May", 6: "Iyun", 7: "Iyul", 8: "Avgust",
    9: "Sentabr", 10: "Oktabr", 11: "Noyabr", 12: "Dekabr"
}


def _set_row_color(row, hex_color: str):
    """Закрашивает строку таблицы в указанный цвет."""
    if not hex_color:
        return
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    trPr.append(shd)


def _add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    return p


def _add_table(doc, headers: list, rows: list, col_widths=None):
    """Добавляет таблицу с заголовком и данными."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Заголовок
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = "Arial"
        cell.paragraphs[0].runs[0].font.size = Pt(11)

    # Данные
    for row_data in rows:
        row = table.add_row()
        color = row_data.get("_color")
        for i, key in enumerate(headers):
            val = row_data.get(key, row_data.get(str(i), ""))
            cell = row.cells[i]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Arial"
                run.font.size = Pt(10)
        if color and color in COLORS:
            _set_row_color(row, COLORS[color])

    return table


def _parse_meeting_json(transcript: str) -> dict:
    """
    Отправляет транскрипт в GPT-4.1 и получает структурированный JSON собрания.
    """
    now = datetime.now()
    next_day = now + timedelta(days=1)

    system_prompt = f"""
Ты — профессиональный секретарь компании Pulinform (взыскание долгов, Узбекистан).
Проанализируй транскрипт ежедневного собрания и верни ТОЛЬКО валидный JSON без markdown.

Текущая дата: {now.strftime('%Y-%m-%d')}
Следующий день: {next_day.strftime('%Y-%m-%d')}
Язык: узбекский латиница

Формат JSON:
{{
  "meeting_number": "порядковый номер собрания или 'N'",
  "date": "{now.strftime('%Y-%m-%d')}",
  "moderator": "имя модератора",
  "participants": [
    {{"name": "имя", "role": "должность", "status": "✅ Ishtirok etdi"}}
  ],
  "sections": [
    {{
      "title": "название раздела",
      "rows": [
        {{
          "col1": "текст", "col2": "текст", "_color": "red|yellow|green|white"
        }}
      ],
      "headers": ["Holat", "Tafsilot"]
    }}
  ],
  "decisions": [
    {{"code": "Q-001", "text": "текст решения", "responsible": "ответственный"}}
  ],
  "tasks": [
    {{"code": "T-001", "text": "задача", "responsible": "ответственный", "deadline": "срок"}}
  ],
  "next_meeting": "{next_day.strftime('%Y-%m-%d')} 09:30",
  "priorities_for_agenda": "краткое описание главных приоритетов следующего дня"
}}

Цветовая логика строк:
- red (F8D7DA): критично, нарушение, проблема
- yellow (FFF3CD): внимание, требует контроля
- green (D4EDDA): выполнено, хороший результат
- white: нейтральная информация
"""

    response = client.chat.completions.create(
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Транскрипт собрания:\n\n{transcript}"},
        ],
        temperature=0.3,
        max_tokens=4000,
    )

    raw = response.choices[0].message.content.strip()
    # Убираем markdown-обёртку если есть
    raw = re.sub(r"^```json\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    return json.loads(raw)


def generate_bayonnoma(data: dict) -> str:
    """Генерирует Bayonnoma.docx и возвращает путь."""
    dt = datetime.strptime(data["date"], "%Y-%m-%d")
    month_uz = MONTHS_UZ[dt.month]
    filename = f"{dt.day:02d}_{month_uz}_{dt.year}_Bayonnoma.docx"
    out_path = str(config.DOCS_DIR / filename)

    doc = Document()

    # Шапка
    doc.add_heading("PULINFORM", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("KUNLIK UCHRASHUV BAYONNOMASI", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        f"Sana: {dt.day}-{month_uz.lower()} {dt.year}  |  "
        f"Yig'ilish №{data.get('meeting_number', 'N')}  |  "
        f"Moderator: {data.get('moderator', 'Begzod')}"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    # Участники
    _add_heading(doc, "ISHTIROKCHILAR", 1)
    participants = data.get("participants", [])
    if participants:
        _add_table(
            doc,
            ["Ism / Viloyat", "Lavozim", "Holat"],
            [{"Ism / Viloyat": p["name"], "Lavozim": p["role"], "Holat": p["status"]}
             for p in participants]
        )
    doc.add_paragraph("")

    # Разделы
    for i, section in enumerate(data.get("sections", []), 1):
        _add_heading(doc, f"{i}. {section['title'].upper()}", 1)
        headers = section.get("headers", ["Masala", "Tafsilot"])
        rows = section.get("rows", [])
        if rows:
            table_rows = []
            for row in rows:
                r = {h: row.get(f"col{j+1}", row.get(h, ""))
                     for j, h in enumerate(headers)}
                r["_color"] = row.get("_color", "white")
                table_rows.append(r)
            _add_table(doc, headers, table_rows)
        doc.add_paragraph("")

    # Реестр решений
    decisions = data.get("decisions", [])
    if decisions:
        _add_heading(doc, "QARORLAR REESTRI", 1)
        _add_table(
            doc,
            ["Kod", "Qaror mazmuni", "Masul"],
            [{"Kod": d["code"], "Qaror mazmuni": d["text"], "Masul": d["responsible"]}
             for d in decisions]
        )
        doc.add_paragraph("")

    # Реестр задач
    tasks = data.get("tasks", [])
    if tasks:
        _add_heading(doc, "TOPSHIRIQLAR REESTRI", 1)
        _add_table(
            doc,
            ["Kod", "Topshiriq mazmuni", "Masul", "Muddat"],
            [{"Kod": t["code"], "Topshiriq mazmuni": t["text"],
              "Masul": t["responsible"], "Muddat": t["deadline"]}
             for t in tasks]
        )
        doc.add_paragraph("")

    # Футер
    next_dt = datetime.strptime(data["next_meeting"].split()[0], "%Y-%m-%d")
    next_month = MONTHS_UZ[next_dt.month]
    footer = doc.add_paragraph(
        f"Keyingi uchrashuv: {next_dt.day}-{next_month.lower()} {next_dt.year}, "
        f"{data['next_meeting'].split()[-1]}  |  "
        "Bayonnomani guruhga yuklash: Begzod  |  AI tomonidan tayyorlangan — "
        "mas'ul shaxs review qilishi shart"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    return out_path


def generate_agenda(data: dict) -> str:
    """Генерирует Agenda.docx для следующего дня и возвращает путь."""
    next_meeting_str = data.get("next_meeting", "")
    try:
        next_dt = datetime.strptime(next_meeting_str.split()[0], "%Y-%m-%d")
    except Exception:
        next_dt = datetime.now() + timedelta(days=1)

    month_uz = MONTHS_UZ[next_dt.month]
    filename = f"{next_dt.day:02d}_{month_uz}_{next_dt.year}_Agenda.docx"
    out_path = str(config.DOCS_DIR / filename)

    # Номер следующего собрания
    try:
        num = int(data.get("meeting_number", "0")) + 1
    except Exception:
        num = "N+1"

    doc = Document()

    # Шапка
    doc.add_heading("PULINFORM", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("KUNLIK UCHRASHUV AGENDASI", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(
        f"{next_dt.day}-{month_uz.lower()} {next_dt.year}, soat "
        f"{next_meeting_str.split()[-1] if next_meeting_str else '09:30'}  |  "
        f"Yig'ilish №{num}"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Красная строка — главные приоритеты
    priorities = data.get("priorities_for_agenda", "Avvalgi kun topshiriqlari nazorati")
    p2 = doc.add_paragraph(priorities)
    p2.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p2.runs[0].font.bold = True

    doc.add_paragraph("")

    # Блоки из задач предыдущего дня
    tasks = data.get("tasks", [])
    decisions = data.get("decisions", [])

    if tasks:
        _add_heading(doc, "TOPSHIRIQLAR NAZORATI", 1)
        _add_table(
            doc,
            ["Kod", "Masala", "Ijrochi", "Bog'lanma", "Prioritet"],
            [{"Kod": f"A-{i+1:03d}", "Masala": t["text"],
              "Ijrochi": t["responsible"], "Bog'lanma": t["code"], "Prioritet": "YUQORI"}
             for i, t in enumerate(tasks)]
        )
        doc.add_paragraph("")

    # Итоговая таблица
    yuqori = len(tasks)
    _add_heading(doc, "UMUMIY", 1)
    _add_table(
        doc,
        ["Jami savol", "YUQORI prioritet", "Faol bo'limlar"],
        [{"Jami savol": str(yuqori), "YUQORI prioritet": str(yuqori),
          "Faol bo'limlar": "Topshiriqlar, Nazorat"}]
    )
    doc.add_paragraph("")

    # Футер
    curr_dt = datetime.strptime(data["date"], "%Y-%m-%d")
    curr_month = MONTHS_UZ[curr_dt.month]
    footer = doc.add_paragraph(
        f"Keyingi uchrashuv: {next_dt.day}-{month_uz.lower()} {next_dt.year}  |  "
        f"Qaror (Q-xxx) va topshiriq (T-xxx) kodlari "
        f"{curr_dt.day}.{curr_dt.month:02d}.{curr_dt.year} bayonnomasiga bog'langan."
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    return out_path


def process_meeting_transcript(transcript: str) -> tuple:
    """
    Главная функция: транскрипт → (bayonnoma_path, agenda_path, summary_text).
    """
    data = _parse_meeting_json(transcript)
    bayonnoma_path = generate_bayonnoma(data)
    agenda_path = generate_agenda(data)

    # Краткое резюме для сообщения в чат
    decisions = data.get("decisions", [])
    tasks = data.get("tasks", [])

    summary_lines = [
        f"📋 *Bayonnoma tayyor* | {data.get('date', '')} | Yig'ilish №{data.get('meeting_number', 'N')}",
        "",
        f"👥 Ishtirokchilar: {len(data.get('participants', []))} nafar",
        f"✅ Qarorlar: {len(decisions)} ta",
        f"📌 Topshiriqlar: {len(tasks)} ta",
        "",
    ]

    if decisions:
        summary_lines.append("*Asosiy qarorlar:*")
        for d in decisions[:3]:
            summary_lines.append(f"• {d['code']}: {d['text'][:80]}")
        if len(decisions) > 3:
            summary_lines.append(f"  ...va yana {len(decisions)-3} ta")

    if tasks:
        summary_lines.append("")
        summary_lines.append("*Bugungi topshiriqlar:*")
        for t in tasks[:4]:
            summary_lines.append(f"• {t['code']} [{t['deadline']}]: {t['text'][:60]}")

    return bayonnoma_path, agenda_path, "\n".join(summary_lines)
